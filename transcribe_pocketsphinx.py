import speech_recognition as sr
from docx import Document
import os

# Initialize the recognizer
recognizer = sr.Recognizer()

# Set the recognizer to use PocketSphinx as the engine
recognizer.recognize_sphinx

# Set the folder path with WAV files
folder_path = "filepath"

# Initialize a list to store transcriptions
transcriptions = []

# Function to format a timestamp in the hh:mm:ss format
def format_timestamp(seconds):
    minutes, seconds = divmod(seconds, 60)
    hours, minutes = divmod(minutes, 60)
    return f"{hours:02}:{minutes:02}:{seconds:02}"

# Create a Word document to store all transcriptions
doc = Document()

# Get a sorted list of WAV files based on numeric order
wav_files = sorted(
    [file for file in os.listdir(folder_path) if file.endswith(".wav")],
    key=lambda x: int(x.split("_")[1].split(".")[0])
)

# Loop through the sorted files
for file in wav_files:
    audio_file_path = os.path.join(folder_path, file)

    # Load the audio file
    with sr.AudioFile(audio_file_path) as source:
        audio_data = recognizer.record(source)

    # Perform speech recognition
    try:
        transcription = recognizer.recognize_sphinx(audio_data)
        print(f"Transcription for {file}: {transcription}")
    except sr.UnknownValueError:
        print(f"Sphinx could not understand the audio in {file}")
        transcription = ""
    except sr.RequestError as e:
        print(f"Sphinx error for {file}: {e}")
        transcription = ""

    # Calculate the audio duration in seconds
    audio_duration = len(audio_data.frame_data) / audio_data.sample_rate

    # Add timestamps every 30 seconds
    timestamp_interval = 30  # Interval in seconds
    timestamp = 0

    while timestamp <= audio_duration:
        timestamp_str = format_timestamp(timestamp)
        transcription = transcription.replace(timestamp_str, f"{file} - {timestamp_str}", 1)
        timestamp += timestamp_interval

        # Add timestamp to the transcription
        transcription += f"\n{file} - {timestamp_str}"

    # Append the transcription to the Word document
    doc.add_heading(f'Transcription for {file}', 0)
    doc.add_paragraph(transcription)

# Save the DOC file with all transcriptions
doc.save("filename")
